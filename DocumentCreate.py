
from docx import *
from docx import Document
from docx.shared import Inches
from docx.text.paragraph import *
from Utilities.GetFunctional import getfunctional
from Utilities.GetRegression import getregression
from Utilities.GetReleaseInfo import getreleaseinfo
from Utilities.GetSummary import getsummary
from Utilities.GetPerformance import getperformance
from Utilities.GetUnit import getunit


def create(id):

    # Create Def Vars to build the document...
    setid = id
    summarydata = getsummary(setid)
    performancedata = getperformance(setid)
    functionaldata = getfunctional(setid)
    regressiondata = getregression(setid)
    unitdata = getunit(setid)
   # text1 = '''"perfect paragraph" will start with a topic sentence. It will have detail sentences in the middle and end with a concluding sentence. It will only cover one topic from start to finish. The length of a paragraph is supposed to be determined by the topic, but often writers will create a paragraph simply to ensure they're not presenting too much text in one chunk.'''
    records = getreleaseinfo(setid)
    text2 = "Content Goes Here!"
    path = "/Users/richardson/python/"
    # Create document
    document = Document()

    document.add_heading("Release Artifacts: " + str(id), 0)

 # Summary Section
    document.add_heading('Release Information', level=1)

    RNDetails = document.add_heading('Release Details', level=2)
    RNDetails.paragraph_format.left_indent = Inches(0.15)

    table1 = document.add_table(
        rows=0, cols=2, style="Table Grid", )
    for cat, rect in records:
        row_cells = table1.add_row().cells
        row_cells[0].paragraphs[0].add_run(cat).bold = True
        row_cells[1].text = str(rect)

    RNNotes = document.add_heading('Release Contents', level=2)
    RNNotes.paragraph_format.left_indent = Inches(0.15)
    RNNotestext = document.add_paragraph(summarydata)
    RNNotestext.paragraph_format.left_indent = Inches(0.25)

   ######################################
   # Release Content Sections
   ######################################
    document.add_heading('Release Contents', level=1)

    RNFeatures = document.add_heading('Release Features: Whats New', level=2)
    RNFeatures.paragraph_format.left_indent = Inches(0.15)
    RNFeaturestext = document.add_paragraph(text2)
    RNFeaturestext.paragraph_format.left_indent = Inches(0.25)

    RNBugs = document.add_heading(
        'Bug Fixes: Anyone need this empty can of RAID?', level=2)
    RNBugs.paragraph_format.left_indent = Inches(0.15)
    RNbugstext = document.add_paragraph(text2)
    RNbugstext.paragraph_format.left_indent = Inches(0.25)

    RNSecurity = document.add_heading(
        'Security Patches: Secure The Hatches! ', level=2)
    RNSecurity.paragraph_format.left_indent = Inches(0.15)
    RNSecuritytext = document.add_paragraph(text2)
    RNSecuritytext.paragraph_format.left_indent = Inches(0.25)

   ########################################
   # Release Quility Gates
   ########################################
    document.add_heading('Release Quality Gates ', level=1)

    QGUnit = document.add_heading('Unit Testing Gate', level=2)
    QGUnit.paragraph_format.left_indent = Inches(0.15)
    QGUnittext = document.add_paragraph(unitdata)
    QGUnittext.paragraph_format.left_indent = Inches(0.25)

    QGFunctional = document.add_heading('Functional Test Gate', level=2)
    QGFunctional.paragraph_format.left_indent = Inches(0.15)
    QGFunctionaltext = document.add_paragraph(functionaldata)
    QGFunctionaltext.paragraph_format.left_indent = Inches(0.25)

    QGRegression = document.add_heading('Regression Test Gate', level=2)
    QGRegression.paragraph_format.left_indent = Inches(0.15)
    QGRegressiontext = document.add_paragraph(regressiondata)
    QGRegressiontext.paragraph_format.left_indent = Inches(0.25)

    QGPerformance = document.add_heading('Performance Test Gate', level=2)
    QGPerformance.paragraph_format.left_indent = Inches(0.15)
    QGPerformancetext = document.add_paragraph(performancedata)
    QGPerformancetext.paragraph_format.left_indent = Inches(0.25)
    document.add_page_break()

    document.save(path + str(setid) + '.docx')
    return
